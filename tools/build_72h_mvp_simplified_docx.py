from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "实训-知勤打卡系统-72小时MVP需求精简与里程碑报告.docx"


def add_field(paragraph, field):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, text, end])


def setup_page(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.4)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def set_header_footer(section, show_footer=True):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.text = ""
    r = header.add_run("实训-知勤打卡系统-72小时MVP")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(9)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = ""
    if not show_footer:
        return
    r = footer.add_run("第 ")
    r.font.name = "宋体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体"); r.font.size = Pt(9)
    add_field(footer, "PAGE")
    r = footer.add_run(" 页 共 ")
    r.font.name = "宋体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体"); r.font.size = Pt(9)
    add_field(footer, "NUMPAGES")
    r = footer.add_run(" 页")
    r.font.name = "宋体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体"); r.font.size = Pt(9)


def patch_page_start(section, start=1):
    sect_pr = section._sectPr
    pg = sect_pr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sect_pr.append(pg)
    pg.set(qn("w:start"), str(start))


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    body = styles["Body Text"]
    body.font.name = "宋体"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    body.font.size = Pt(12)
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.first_line_indent = Pt(21)
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(0)

    for style_name, size, align in [
        ("Heading 1", 16, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", 15, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 14, WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        st = styles[style_name]
        st.font.name = "黑体"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.font.size = Pt(size)
        st.font.bold = False
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.alignment = align
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(0)

    cap = styles["Caption"]
    cap.font.name = "宋体"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cap.font.size = Pt(10.5)
    cap.font.italic = False
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def heading(doc, text, level=1, page_break=True):
    if level == 1:
        if page_break:
            doc.add_page_break()
        p = doc.add_paragraph(style="Heading 1")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        p = doc.add_paragraph(style="Heading 2")
    else:
        p = doc.add_paragraph(style="Heading 3")
    p.add_run(text)


def body(doc, text):
    p = doc.add_paragraph(style="Body Text")
    p.paragraph_format.first_line_indent = Pt(21)
    p.add_run(text)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Body Text")
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Pt(21)
        p.paragraph_format.hanging_indent = Pt(10.5)
        p.add_run("• " + item)


def cell_text(cell, text, bold=False, size=10.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    r = p.add_run(str(text))
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(size)
    r.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def width(cell, cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(cm).twips)))
    tc_w.set(qn("w:type"), "dxa")


def table(doc, caption, headers, rows, widths):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(caption)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        cell_text(c, h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(c, "E8EEF5")
        width(c, widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value)
            width(cells[i], widths[i])
    doc.add_paragraph()


def build():
    doc = Document()
    setup_page(doc.sections[0])
    setup_styles(doc)
    set_header_footer(doc.sections[0], show_footer=False)

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("实训-知勤打卡系统\n72小时MVP需求精简与里程碑报告")
    r.font.name = "黑体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(23); r.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("高中生可读版 · Python + FastAPI + Vue + MySQL")
    r.font.name = "宋体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(14)
    for _ in range(5):
        doc.add_paragraph()
    for line in [
        "项目名称：知勤打卡系统（AI 思政打卡 MVP）",
        "文档版本：V1.0",
        "项目周期：72 小时极限 MVP",
        "技术约束：Python + FastAPI + Vue + MySQL",
        "编制日期：2026 年 6 月 23 日",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = "宋体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(10.5)

    doc.add_section(WD_SECTION.NEW_PAGE)
    setup_page(doc.sections[-1])
    set_header_footer(doc.sections[-1], show_footer=True)
    patch_page_start(doc.sections[-1], 1)

    heading(doc, "1 一句话说明这个项目", 1, page_break=False)
    body(doc, "知勤打卡系统可以理解成“学校版打卡小程序”。学生每天在手机上完成查寝、上课或实习打卡；老师在后台看谁打了、谁没打、谁需要审核；管理员看整体数据。")
    body(doc, "如果把学校管理比作班主任点名，传统方式是拿纸笔一个个问；这个系统就是把点名、记录、审核、统计放到一个数字工具里。72 小时 MVP 不做大而全，只做最短可跑通的主路。")

    heading(doc, "2 为什么要精简需求", 1)
    body(doc, "72 小时只有 3 天，不能把所有想法都做进去。产品经理最重要的判断是：先做“没有它系统就不能用”的功能，其他好看的、复杂的、锦上添花的功能先放到下一版。")
    table(doc, "表 2.1 需求精简原则", ["原则", "高中生能懂的解释", "本项目怎么做"], [
        ["先闭环", "先让事情从开始到结束能走通", "学生打卡 → 老师审核 → 学生看到结果 → 统计更新"],
        ["少依赖", "别把成败绑在不确定的外部服务上", "AI 审核、人脸识别先不做强制功能"],
        ["可演示", "老师一看就知道系统有用", "准备演示账号和演示数据"],
        ["可扩展", "先搭骨架，以后再加肉", "数据库预留状态字段，接口按模块拆开"],
    ], [2.4, 5.4, 6.0])

    heading(doc, "3 72 小时 MVP 只做什么", 1)
    table(doc, "表 3.1 MVP 功能范围", ["优先级", "功能", "是否进入 72 小时", "原因"], [
        ["Must", "登录与角色识别", "进入", "不知道谁登录，就无法区分学生、老师和管理员"],
        ["Must", "学生三类打卡", "进入", "这是系统的核心，没有打卡就不是打卡系统"],
        ["Must", "打卡历史与审核结果", "进入", "学生需要知道自己是否完成任务"],
        ["Must", "教师审核", "进入", "老师必须能处理待审核和异常记录"],
        ["Must", "统计概览", "进入", "演示时要看到管理价值"],
        ["Should", "定位记录", "简化进入", "只记录经纬度和异常，不做复杂范围校验"],
        ["Could", "AI 内容审核", "暂不进入", "外部接口不稳定，可能拖慢 72 小时交付"],
        ["Could", "人脸识别", "暂不进入", "SDK、隐私和准确率都需要额外验证"],
        ["Could", "补签、推送、积分", "暂不进入", "有价值，但不是第一条主路"],
    ], [2.0, 3.6, 3.0, 5.2])

    heading(doc, "4 用户和他们要完成的事", 1)
    table(doc, "表 4.1 用户角色", ["角色", "最想做什么", "系统给他的入口"], [
        ["学生", "快速完成今天的打卡，知道老师有没有通过", "学生端：今日状态、提交打卡、历史记录"],
        ["辅导员/教师", "快速看到班级谁打了、谁没打，并完成审核", "教师端：审核列表、详情、通过/驳回"],
        ["管理员", "看全校或整体数据，确认系统运行情况", "管理端：统计概览、基础数据"],
    ], [2.4, 6.0, 5.4])

    heading(doc, "5 用高中生能懂的方式看技术架构", 1)
    body(doc, "这个系统可以想成一家餐厅：前端 Vue 是菜单和点餐界面，后端 FastAPI 是厨房，MySQL 是仓库，Git 是大家一起写菜谱时用的版本记录本。")
    table(doc, "表 5.1 技术分工解释", ["技术", "像什么", "负责什么", "2026 最佳实践"], [
        ["Vue", "菜单和点餐屏幕", "学生和老师看到的页面", "页面拆成组件，重复逻辑放 composables"],
        ["FastAPI", "厨房和服务员", "接收请求、检查权限、处理业务", "用 APIRouter 按模块拆路由"],
        ["MySQL", "仓库和账本", "保存用户、打卡、审核、统计数据", "用 InnoDB、外键、索引保证数据清楚"],
        ["Git/GitHub", "多人协作的版本记录本", "记录谁改了什么，避免互相覆盖", "main 分支保护，功能走 Pull Request"],
    ], [2.4, 3.2, 4.2, 5.2])

    heading(doc, "6 推荐模块拆分", 1)
    table(doc, "表 6.1 MVP 模块清单", ["模块", "主要功能", "负责人建议", "独立开发边界"], [
        ["认证模块", "登录、当前用户、角色判断", "后端负责人", "只提供 /auth 接口，前端只调用接口"],
        ["学生打卡模块", "今日状态、提交打卡、历史记录", "学生端 + 后端", "学生端可先用 mock 数据开发"],
        ["教师审核模块", "审核列表、详情、通过/驳回", "教师端 + 后端", "教师端不依赖学生端页面"],
        ["统计模块", "打卡率、待审核数、异常数", "后端 + 管理端", "只读接口，不阻塞打卡主流程"],
        ["数据库模块", "用户、班级、打卡、审核表", "后端负责人", "前端不能直接连数据库"],
        ["验收模块", "演示脚本、缺陷清单、四环境检查", "产品/QA", "用文档和清单管理，不改业务代码"],
    ], [2.8, 4.8, 3.0, 5.0])

    heading(doc, "7 72 小时里程碑设计", 1)
    body(doc, "里程碑就是“每隔一段时间检查一次有没有跑偏”。72 小时开发最怕最后一天才发现前后端接不上，所以必须每 12 小时做一次可运行检查。")
    table(doc, "表 7.1 72 小时里程碑", ["时间", "目标", "必须看到什么", "如果失败怎么办"], [
        ["H0-H6", "需求冻结", "大家确认只做 Must 功能", "砍掉所有非主路功能"],
        ["H6-H12", "项目骨架", "FastAPI、两个 Vue 前端、MySQL 能启动", "先用最简单目录和假数据"],
        ["H12-H24", "登录和学生打卡", "学生能登录并提交一条打卡", "定位、照片都可先简化"],
        ["H24-H36", "教师审核", "老师能看到打卡并通过/驳回", "先做列表和按钮，后补美化"],
        ["H36-H48", "统计概览", "能看到打卡率、待审核数、异常数", "趋势图先用占位"],
        ["H48-H60", "全流程联调", "学生打卡、老师审核、学生看结果都通", "只修主流程 P0/P1 问题"],
        ["H60-H72", "验收和部署", "本地、服务器、UAT、生产测试完成记录", "不能上线就记录原因并本地演示"],
    ], [2.0, 3.4, 5.5, 4.5])

    heading(doc, "8 每个阶段的交付物", 1)
    table(doc, "表 8.1 阶段交付物", ["阶段", "交付物", "谁检查"], [
        ["需求冻结", "MVP 功能清单、接口草案、验收脚本", "产品经理和全员"],
        ["项目骨架", "后端 /health、前端首页、数据库连接", "技术负责人"],
        ["学生打卡", "登录页、今日状态页、提交打卡接口", "学生端负责人和后端负责人"],
        ["教师审核", "审核列表、详情、通过/驳回接口", "教师端负责人和后端负责人"],
        ["统计概览", "统计接口、指标卡片、异常入口", "管理端负责人"],
        ["最终验收", "演示账号、演示数据、验收记录", "产品/QA"],
    ], [3.0, 7.0, 4.0])

    heading(doc, "9 最小数据库设计", 1)
    body(doc, "数据库设计不要把所有东西塞进一张表。像整理书包一样，课本、作业本、文具要分开放；系统里用户、打卡、审核也要分开放。")
    table(doc, "表 9.1 最小数据表", ["表名", "保存什么", "为什么需要"], [
        ["users", "学生、教师、管理员账号", "登录和权限判断"],
        ["classes", "班级信息", "老师按班级查看学生"],
        ["checkins", "学生打卡记录", "系统最核心的数据"],
        ["reviews", "教师审核结果", "记录谁审核、通过还是驳回"],
        ["courses", "课程和班级关系", "上课打卡后续扩展"],
        ["location_rules", "地点规则", "后续做定位范围校验"],
    ], [3.0, 5.0, 5.5])

    heading(doc, "10 最小接口清单", 1)
    table(doc, "表 10.1 API 清单", ["模块", "接口", "给谁用", "作用"], [
        ["认证", "POST /api/v1/auth/login", "学生端、教师端", "登录"],
        ["认证", "GET /api/v1/auth/me", "学生端、教师端", "获取当前用户"],
        ["打卡", "GET /api/v1/checkins/today", "学生端", "获取今日状态"],
        ["打卡", "POST /api/v1/checkins", "学生端", "提交打卡"],
        ["打卡", "GET /api/v1/checkins", "学生端", "查看历史"],
        ["审核", "GET /api/v1/teacher/checkins", "教师端", "查看班级打卡"],
        ["审核", "POST /api/v1/teacher/reviews", "教师端", "提交审核结果"],
        ["统计", "GET /api/v1/stats/overview", "管理端", "查看统计概览"],
    ], [2.4, 5.0, 3.4, 4.8])

    heading(doc, "11 最终验收脚本", 1)
    table(doc, "表 11.1 10 步验收脚本", ["步骤", "动作", "通过标准"], [
        ["1", "学生账号登录", "进入学生首页"],
        ["2", "学生查看今日三类状态", "查寝、上课、实习状态可见"],
        ["3", "学生提交查寝打卡", "提交成功，状态为待审核"],
        ["4", "学生重复提交查寝", "系统提示今日已打卡"],
        ["5", "教师账号登录", "进入教师审核工作台"],
        ["6", "教师筛选今日待审核", "能看到学生打卡记录"],
        ["7", "教师打开详情", "能看到内容、定位、时间"],
        ["8", "教师审核通过或驳回", "状态被更新"],
        ["9", "学生查看历史", "能看到审核结果"],
        ["10", "查看统计概览", "打卡率、待审核数、异常数可见"],
    ], [1.4, 5.5, 7.5])

    heading(doc, "12 风险和降级方案", 1)
    table(doc, "表 12.1 风险清单", ["风险", "影响", "降级方案"], [
        ["AI 接口不稳定", "拖慢提交打卡", "先不接 AI，状态统一待教师审核"],
        ["定位不准或被浏览器拒绝", "学生无法提交", "允许提交，标记定位 unknown"],
        ["页面做不完", "演示不完整", "先做核心按钮和列表，少做动画和美化"],
        ["部署失败", "四环境验收受影响", "先保证本地和服务器二选一可演示，并记录问题"],
        ["接口字段变来变去", "前后端联调失败", "先写 API 契约，字段冻结后再开发"],
    ], [4.0, 4.5, 7.0])

    heading(doc, "13 参考资料", 1)
    bullets(doc, [
        "FastAPI 官方文档 Bigger Applications：APIRouter 可以像小型 FastAPI 应用一样组织多文件路由。https://fastapi.tiangolo.com/tutorial/bigger-applications/",
        "Vue 官方文档 Composables：composable 用来封装和复用有状态逻辑。https://vuejs.org/guide/reusability/composables",
        "MySQL 8.4 官方文档 Foreign Key Constraints：InnoDB 支持外键约束并用于维护引用完整性。https://dev.mysql.com/doc/refman/8.4/en/create-table-foreign-keys.html",
        "GitHub Docs Protected Branches：分支保护可限制删除、强推，并要求状态检查或 Pull Request。https://docs.github.com/repositories/configuring-branches-and-merges-in-your-repository/managing-protected-branches/about-protected-branches",
        "GitHub Docs Pull Request Reviews：PR Review 可以帮助团队在合并前评论、建议和审批代码。https://docs.github.com/pull-requests/collaborating-with-pull-requests/reviewing-changes-in-pull-requests/about-pull-request-reviews",
    ])

    doc.core_properties.title = "实训-知勤打卡系统-72小时MVP需求精简与里程碑报告"
    doc.core_properties.author = "ZeKin 项目组"
    doc.core_properties.subject = "72小时MVP、需求精简、里程碑设计"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
