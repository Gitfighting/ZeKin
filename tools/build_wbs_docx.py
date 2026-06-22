from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "实训-知勤打卡系统-WBS任务分解报告.docx"


def add_field(paragraph, field):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def patch_page_number_type(section, fmt="decimal", start=None):
    sect_pr = section._sectPr
    pg = sect_pr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sect_pr.append(pg)
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))


def setup_section(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.4)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def set_header_footer(section, show_footer=True, literal_footer=None):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.text = ""
    r = header.add_run("实训-知勤打卡系统-WBS任务分解")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(9)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = ""
    if not show_footer:
        return
    if literal_footer:
        rr = footer.add_run(literal_footer)
        rr.font.name = "宋体"
        rr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        rr.font.size = Pt(9)
        return
    rr = footer.add_run("第 ")
    rr.font.name = "宋体"; rr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体"); rr.font.size = Pt(9)
    add_field(footer, "PAGE")
    rr = footer.add_run(" 页 共 ")
    rr.font.name = "宋体"; rr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体"); rr.font.size = Pt(9)
    add_field(footer, "NUMPAGES")
    rr = footer.add_run(" 页")
    rr.font.name = "宋体"; rr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体"); rr.font.size = Pt(9)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    body = styles["Body Text"]
    body.font.name = "宋体"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    body.font.size = Pt(12)
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(0)
    body.paragraph_format.first_line_indent = Pt(21)

    for style_name, size, align in [
        ("Heading 1", 16, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", 15, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 14, WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        st = styles[style_name]
        st.font.name = "黑体"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.font.size = Pt(size)
        st.font.bold = False
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.alignment = align
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(0)

    cap = styles["Caption"]
    cap.font.name = "宋体"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cap.font.size = Pt(10.5)
    cap.font.italic = False
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing = 1.5

    code = styles.add_style("Code", 1)
    code.font.name = "Times New Roman"
    code.font.size = Pt(12)
    code.paragraph_format.line_spacing = 1.5


def add_heading(doc, text, level=1, page_break=True):
    if level == 1:
        if page_break:
            doc.add_page_break()
        p = doc.add_paragraph(style="Heading 1")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        p = doc.add_paragraph(style="Heading 2")
    else:
        p = doc.add_paragraph(style="Heading 3")
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(style="Body Text")
    p.paragraph_format.first_line_indent = Pt(21)
    p.add_run(text)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Body Text")
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Pt(21)
        p.paragraph_format.hanging_indent = Pt(10.5)
        p.add_run("• " + item)


def set_cell_text(cell, text, bold=False, size=10.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    r = p.add_run(str(text))
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(size)
    r.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))
    tc_w.set(qn("w:type"), "dxa")


def add_table(doc, caption, headers, rows, widths):
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
        set_cell_width(table.rows[0].cells[i], widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
            set_cell_width(cells[i], widths[i])
    doc.add_paragraph()


def add_code(doc, code):
    for line in code.splitlines():
        p = doc.add_paragraph(style="Code")
        p.add_run(line if line else " ")


def build():
    doc = Document()
    setup_section(doc.sections[0])
    setup_styles(doc)
    set_header_footer(doc.sections[0], show_footer=False)

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("实训-知勤打卡系统\nWBS 任务分解报告")
    r.font.name = "黑体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(24); r.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("模块划分 · 个人分工 · Git 协作 · 72 小时极限开发")
    r.font.name = "宋体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(14)
    for _ in range(5):
        doc.add_paragraph()
    for line in [
        "项目名称：知勤打卡系统（ZeKin）",
        "文档版本：V1.0",
        "技术约束：Python + FastAPI + Vue + MySQL",
        "团队成员：王韩韵、赵耀、胡钊炫、华心仪",
        "编制日期：2026 年 6 月 22 日",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = "宋体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(10.5)

    doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(doc.sections[-1])
    patch_page_number_type(doc.sections[-1], "lowerRoman", 1)
    set_header_footer(doc.sections[-1], show_footer=True, literal_footer="第 i 页 共 14 页")
    h = doc.add_paragraph(style="Heading 1"); h.alignment = WD_ALIGN_PARAGRAPH.CENTER; h.add_run("摘  要")
    add_body(doc, "本文档将知勤打卡系统 72 小时 MVP 需求转化为可执行的工程开发计划，围绕 Python、FastAPI、Vue、MySQL 技术约束，明确模块边界、个人职责、WBS 任务分解、Git 分支管理、接口契约与验收标准。")
    add_body(doc, "本报告强调独立开发、互不干扰：王韩韵负责产品规格与验收，赵耀负责后端 API、数据库与部署，胡钊炫负责教师/管理端，华心仪负责学生端。所有跨模块协作均通过接口契约和 Pull Request 完成。")

    doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(doc.sections[-1])
    set_header_footer(doc.sections[-1], show_footer=False)
    h = doc.add_paragraph(style="Heading 1"); h.alignment = WD_ALIGN_PARAGRAPH.CENTER; h.add_run("目  录")
    toc = [
        ("1 报告目标", "1"),
        ("2 技术最佳实践约束", "2"),
        ("3 总体模块架构", "3"),
        ("4 独立开发边界", "4"),
        ("5 模块拆分说明", "5"),
        ("6 WBS 任务分解", "6"),
        ("7 接口契约与并行开发边界", "9"),
        ("8 Git 版本管理方案", "10"),
        ("9 72 小时里程碑与验收", "10"),
        ("10 风险与完成定义", "11"),
    ]
    for name, page in toc:
        p = doc.add_paragraph()
        r = p.add_run(f"{name}{'.' * max(4, 48 - len(name))}{page}")
        r.font.name = "宋体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(12)

    doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(doc.sections[-1])
    patch_page_number_type(doc.sections[-1], "decimal", 1)
    set_header_footer(doc.sections[-1], show_footer=True)

    add_heading(doc, "1 报告目标", 1, page_break=False)
    add_body(doc, "本报告将需求分析成果转化为工程团队可以直接执行的开发计划，重点解决 72 小时极限开发条件下的模块划分、人员分工、任务依赖、Git 协作和验收标准问题。")
    add_bullets(doc, [
        "明确四位成员的 Owner 目录和职责边界。",
        "将学生打卡、教师审核、统计概览拆解为可并行开发模块。",
        "用 WBS 编号管理任务输入、输出、依赖和验收标准。",
        "通过接口契约和 Git 分支规则避免多人开发互相干扰。",
    ])

    add_heading(doc, "2 技术最佳实践约束", 1)
    add_heading(doc, "2.1 FastAPI 后端约束", 2)
    add_body(doc, "后端采用 FastAPI 模块化结构，使用 APIRouter 拆分认证、打卡、审核、统计等业务路由。router 只负责接收请求和返回响应，业务规则放在 services，数据库访问放在 repositories。")
    add_heading(doc, "2.2 Vue 前端约束", 2)
    add_body(doc, "前端拆分为学生端 H5 与教师/管理端 Web 两个独立应用。学生端和管理端分别维护自己的 views、components、api、stores 和 composables，避免两个同学修改同一套页面造成冲突。")
    add_heading(doc, "2.3 MySQL 数据库约束", 2)
    add_body(doc, "数据库采用 MySQL 8.x 与 InnoDB，使用主键、外键、唯一约束和必要索引保证数据一致性。同一学生同一自然日同一打卡类型只能存在一条有效打卡记录。")
    add_heading(doc, "2.4 Git 协作约束", 2)
    add_body(doc, "Git 使用 main、develop、feature 分支模型。main 保存稳定演示版本，develop 用于 72 小时集成，个人 feature 分支按模块开发。任何跨模块变更必须通过 Pull Request 和 Owner Review。")

    add_heading(doc, "3 总体模块架构", 1)
    add_table(doc, "表 3.1 系统模块总览", ["模块", "负责人", "核心职责"], [
        ["产品规格与验收", "王韩韵", "需求冻结、接口契约、验收脚本、缺陷分级、四环境验收"],
        ["后端 API 与数据库", "赵耀", "FastAPI、MySQL、认证、打卡、审核、统计、部署"],
        ["学生端 H5", "华心仪", "登录、今日状态、提交打卡、历史记录、移动端适配"],
        ["教师/管理端 Web", "胡钊炫", "教师登录、审核列表、详情面板、审核操作、统计概览"],
    ], [4.0, 2.8, 7.2])

    add_heading(doc, "4 独立开发边界", 1)
    add_table(doc, "表 4.1 个人 Owner 目录", ["成员", "Owner 目录", "禁止直接修改"], [
        ["王韩韵", "docs/、acceptance/", "backend/、frontend/ 业务代码"],
        ["赵耀", "backend/、database/、deploy/", "frontend/student/、frontend/admin/ 页面"],
        ["胡钊炫", "frontend/admin/", "backend/ 数据库和业务规则"],
        ["华心仪", "frontend/student/", "backend/ 数据库和业务规则"],
    ], [2.4, 5.8, 5.8])
    add_body(doc, "如果确实需要跨模块修改，必须先在 issue 或群内说明变更原因，并在 Pull Request 中 @ 对应 Owner 审核。")

    add_heading(doc, "5 模块拆分说明", 1)
    modules = [
        ("5.1 M1 认证与角色模块", "让学生、教师、管理员以合法身份进入系统，并进入对应工作台。后端提供登录、当前用户和角色权限；前端实现登录页、登录态保存和未登录拦截。"),
        ("5.2 M2 学生打卡模块", "学生完成查寝、上课、实习三类打卡，系统保存可审核、可追溯记录。学生端可先用 mock 数据开发，后端接口完成后替换真实 API。"),
        ("5.3 M3 教师审核模块", "教师查看所管班级打卡记录并完成通过或驳回。教师端依赖审核接口，不依赖学生端页面完成。"),
        ("5.4 M4 统计概览模块", "教师和管理员查看今日应打卡、已打卡、打卡率、待审核和异常数。该模块只读，不阻塞打卡和审核闭环。"),
        ("5.5 M5 数据库与演示数据模块", "用 users、classes、courses、checkins、reviews、location_rules 六张表支撑 MVP 验收。"),
        ("5.6 M6 验收与项目管理模块", "提供需求跟踪矩阵、验收脚本、演示账号、缺陷模板和四环境检查清单。"),
    ]
    for title, text in modules:
        add_heading(doc, title, 2)
        add_body(doc, text)

    add_heading(doc, "6 WBS 任务分解", 1)
    add_table(doc, "表 6.1 WBS 总览", ["WBS", "工作包", "负责人", "产出物"], [
        ["1.0", "项目规格与验收管理", "王韩韵", "需求冻结、验收脚本、缺陷模板"],
        ["2.0", "后端 API 与数据库", "赵耀", "FastAPI 服务、MySQL 表、接口"],
        ["3.0", "学生端 H5", "华心仪", "登录、今日状态、打卡、历史"],
        ["4.0", "教师/管理端 Web", "胡钊炫", "审核、详情、统计"],
        ["5.0", "联调与四环境验收", "全员", "可演示系统、验收记录"],
    ], [1.6, 4.4, 2.2, 6.2])
    add_table(doc, "表 6.2 王韩韵 WBS", ["WBS", "任务", "输出", "时间"], [
        ["1.1", "冻结 72 小时 MVP 范围", "Must/Should/Could 清单", "H0-H2"],
        ["1.2", "编写接口契约草案", "docs/api-contract.md", "H0-H4"],
        ["1.3", "编写验收脚本", "acceptance/uat-script.md", "H2-H6"],
        ["1.4", "建立缺陷分级规则", "bug-template.md", "H4-H6"],
        ["1.5", "组织每 12 小时集成检查", "集成记录", "H12/H24/H36/H48/H60"],
        ["1.6", "组织最终验收", "验收报告", "H66-H72"],
    ], [1.6, 5.2, 4.4, 2.8])
    add_table(doc, "表 6.3 赵耀 WBS", ["WBS", "任务", "输出", "时间"], [
        ["2.1", "初始化 FastAPI 项目", "backend/app/main.py", "H6-H8"],
        ["2.2", "建立分层目录", "api、schemas、services、repositories", "H6-H8"],
        ["2.3", "设计 MySQL DDL", "database/schema.sql", "H8-H10"],
        ["2.4", "初始化演示数据", "database/seed.sql", "H10-H12"],
        ["2.5", "实现认证接口", "/auth/login、/auth/me", "H12-H18"],
        ["2.6", "实现学生打卡接口", "/checkins/*", "H18-H30"],
        ["2.7", "实现教师审核接口", "/teacher/*", "H30-H42"],
        ["2.8", "实现统计接口", "/stats/overview", "H42-H50"],
        ["2.9", "权限和异常处理", "401、403、统一错误响应", "H48-H56"],
        ["2.10", "部署与环境变量", ".env.example、启动说明", "H56-H66"],
    ], [1.5, 4.8, 5.8, 2.4])
    add_table(doc, "表 6.4 华心仪 WBS", ["WBS", "任务", "输出", "时间"], [
        ["3.1", "初始化 Vue 学生端", "frontend/student/", "H6-H8"],
        ["3.2", "登录页", "LoginView.vue", "H8-H14"],
        ["3.3", "路由与登录态", "router、store", "H12-H18"],
        ["3.4", "今日状态页面", "HomeView.vue", "H18-H26"],
        ["3.5", "打卡表单", "CheckinView.vue", "H24-H34"],
        ["3.6", "定位 composable", "useGeolocation", "H26-H34"],
        ["3.7", "历史记录页面", "HistoryView.vue", "H34-H42"],
        ["3.8", "学生端联调", "可演示学生端", "H42-H58"],
        ["3.9", "移动端适配", "CSS 修正", "H58-H66"],
    ], [1.5, 4.8, 5.8, 2.4])
    add_table(doc, "表 6.5 胡钊炫 WBS", ["WBS", "任务", "输出", "时间"], [
        ["4.1", "初始化 Vue 管理端", "frontend/admin/", "H6-H8"],
        ["4.2", "教师/管理员登录页", "LoginView.vue", "H8-H14"],
        ["4.3", "管理端布局", "Layout、导航、路由", "H12-H18"],
        ["4.4", "审核筛选栏", "筛选组件", "H18-H26"],
        ["4.5", "审核列表", "ReviewList.vue", "H24-H34"],
        ["4.6", "打卡详情面板", "ReviewDetail.vue", "H30-H38"],
        ["4.7", "通过/驳回操作", "审核按钮和原因输入", "H36-H46"],
        ["4.8", "统计概览页面", "DashboardView.vue", "H42-H54"],
        ["4.9", "管理端联调", "可演示管理端", "H54-H66"],
    ], [1.5, 4.8, 5.8, 2.4])

    add_heading(doc, "7 接口契约与并行开发边界", 1)
    add_table(doc, "表 7.1 最小接口清单", ["模块", "方法", "路径", "调用方"], [
        ["Auth", "POST", "/api/v1/auth/login", "学生端、管理端"],
        ["Auth", "GET", "/api/v1/auth/me", "学生端、管理端"],
        ["Checkin", "GET", "/api/v1/checkins/today", "学生端"],
        ["Checkin", "POST", "/api/v1/checkins", "学生端"],
        ["Teacher", "GET", "/api/v1/teacher/checkins", "管理端"],
        ["Teacher", "POST", "/api/v1/teacher/reviews", "管理端"],
        ["Stats", "GET", "/api/v1/stats/overview", "管理端"],
    ], [2.0, 2.0, 5.8, 4.2])
    add_body(doc, "接口字段和枚举值以 docs/api-contract.md 为准。前端展示可翻译为中文，但接口传输必须使用英文枚举，避免前后端各自翻译导致联调失败。")

    add_heading(doc, "8 Git 版本管理方案", 1)
    add_table(doc, "表 8.1 分支设计", ["分支", "用途", "合并方向"], [
        ["main", "稳定演示版本", "只接收 develop 稳定合并"],
        ["develop", "72 小时集成版本", "接收 feature 分支"],
        ["feat/backend-mvp", "后端 API、数据库、部署", "合并到 develop"],
        ["feat/student-h5", "学生端 H5", "合并到 develop"],
        ["feat/admin-review", "教师/管理端 Web", "合并到 develop"],
    ], [4.0, 5.0, 5.0])
    add_heading(doc, "8.2 提交规范", 2)
    add_code(doc, "feat(auth): implement login api\nfeat(checkin): add submit checkin service\nfeat(student): build today status view\nfeat(admin): build teacher review list\ndocs(wbs): update 72h task plan\nfix(review): require reject comment")

    add_heading(doc, "9 72 小时里程碑与验收", 1)
    add_table(doc, "表 9.1 72 小时里程碑", ["时间点", "里程碑", "判断标准"], [
        ["H6", "需求和接口冻结", "所有人确认 Must 范围"],
        ["H12", "项目骨架可启动", "后端和两个前端均可启动"],
        ["H24", "学生打卡最小闭环", "学生可登录并提交打卡"],
        ["H36", "教师审核最小闭环", "教师可审核一条记录"],
        ["H48", "统计概览可用", "指标卡片返回真实数据"],
        ["H60", "全流程联调完成", "验收脚本主路径通过"],
        ["H72", "四环境验收", "本地、服务器、UAT、生产测试完成记录"],
    ], [2.2, 4.6, 7.0])
    add_body(doc, "最终验收脚本包括学生登录、查看今日状态、提交查寝、重复提交拦截、教师登录、筛选待审核、审核通过、学生查看审核结果、统计页查看指标等 10 个步骤。")

    add_heading(doc, "10 风险与完成定义", 1)
    add_table(doc, "表 10.1 风险与降级策略", ["风险", "影响", "降级策略"], [
        ["登录鉴权耗时过长", "阻塞联调", "使用简单 Token 或固定演示账号"],
        ["定位 API 权限失败", "学生端打卡受阻", "允许 mock 经纬度并标记 unknown"],
        ["教师审核联调延迟", "管理闭环受阻", "先用 mock 列表完成页面"],
        ["统计接口复杂", "演示受阻", "只返回今日核心指标"],
        ["服务器部署失败", "环境验收受阻", "先用局域网演示并记录缺陷"],
    ], [4.0, 3.6, 6.4])
    add_heading(doc, "10.2 Definition of Done", 2)
    add_bullets(doc, [
        "任务关联 WBS 编号，代码或文档已提交到对应分支。",
        "不修改非 Owner 目录，或跨模块修改已获得 Owner 确认。",
        "涉及接口的任务已同步接口契约。",
        "涉及数据库的任务已提供建表或迁移脚本。",
        "学生打卡、教师审核、统计概览三条核心路径可演示。",
        "GitHub main 分支保存最终可演示版本。",
    ])

    doc.core_properties.title = "实训-知勤打卡系统-WBS任务分解报告"
    doc.core_properties.subject = "模块划分、个人分工与 WBS"
    doc.core_properties.author = "ZeKin 项目组"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
