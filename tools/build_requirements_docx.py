from __future__ import annotations

import os
import re
import textwrap
import zlib
from pathlib import Path
from typing import Iterable

import urllib.request
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "实训-知勤打卡系统-需求分析.docx"
ASSET_DIR = ROOT / "build" / "docx_assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)


def plantuml_encode(text: str) -> str:
    data = zlib.compress(text.encode("utf-8"))[2:-4]
    alphabet = "0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz-_"

    def append3(b1, b2, b3):
        c1 = b1 >> 2
        c2 = ((b1 & 0x3) << 4) | (b2 >> 4)
        c3 = ((b2 & 0xF) << 2) | (b3 >> 6)
        c4 = b3 & 0x3F
        return alphabet[c1] + alphabet[c2] + alphabet[c3] + alphabet[c4]

    res = []
    for i in range(0, len(data), 3):
        chunk = data[i : i + 3]
        if len(chunk) == 3:
            res.append(append3(chunk[0], chunk[1], chunk[2]))
        elif len(chunk) == 2:
            res.append(append3(chunk[0], chunk[1], 0))
        else:
            res.append(append3(chunk[0], 0, 0))
    return "".join(res)


def render_plantuml() -> dict[str, Path]:
    rendered = {}
    for puml in sorted((ROOT / "diagrams").glob("*.puml")):
        text = puml.read_text(encoding="utf-8")
        url = f"https://www.plantuml.com/plantuml/png/{plantuml_encode(text)}"
        out = ASSET_DIR / f"{puml.stem}.png"
        try:
            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=20) as resp:
                content = resp.read()
            if content.startswith(b"\x89PNG"):
                out.write_bytes(content)
                rendered[puml.stem] = out
        except Exception:
            pass
    return rendered


def font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for c in candidates:
        if Path(c).exists():
            return ImageFont.truetype(c, size)
    return ImageFont.load_default()


def draw_round(draw, xy, fill, outline=None, radius=18, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def make_prototype_images() -> dict[str, Path]:
    images = {}

    # Student mobile prototype
    img = Image.new("RGB", (780, 1400), "#f8fafc")
    d = ImageDraw.Draw(img)
    f_title = font(42, True)
    f_h = font(30, True)
    f = font(24)
    f_b = font(25, True)
    d.rectangle([0, 0, 780, 150], fill="#1d4ed8")
    d.text((42, 45), "ZeKin 学生打卡", fill="white", font=f_title)
    d.text((42, 100), "今日任务：查寝 / 上课 / 实习", fill="#bfdbfe", font=f)
    draw_round(d, [36, 190, 744, 400], "#ffffff", "#dbeafe")
    d.text((68, 222), "今日状态", fill="#0f172a", font=f_h)
    for i, (name, state, fill, color) in enumerate(
        [("查寝", "已提交", "#dcfce7", "#166534"), ("上课", "待打卡", "#fef3c7", "#92400e"), ("实习", "未开始", "#e0f2fe", "#075985")]
    ):
        x = 68 + i * 220
        draw_round(d, [x, 280, x + 170, 360], fill, None, 18)
        d.text((x + 26, 296), name, fill=color, font=f)
        d.text((x + 26, 326), state, fill=color, font=f_b)
    d.text((48, 470), "提交打卡", fill="#0f172a", font=f_h)
    draw_round(d, [36, 510, 744, 1060], "#ffffff", "#e2e8f0")
    d.text((68, 550), "打卡类型", fill="#475569", font=f)
    for i, name in enumerate(["查寝", "上课", "实习"]):
        x = 68 + i * 180
        fillc = "#1d4ed8" if i == 0 else "#eff6ff"
        textc = "white" if i == 0 else "#1d4ed8"
        draw_round(d, [x, 590, x + 140, 650], fillc, "#bfdbfe", 30)
        d.text((x + 43, 606), name, fill=textc, font=f_b)
    d.text((68, 720), "思政心得 / 实习日志", fill="#475569", font=f)
    draw_round(d, [68, 760, 712, 940], "#f8fafc", "#cbd5e1", 18)
    d.text((96, 800), "请输入 10-300 字内容...", fill="#94a3b8", font=f)
    draw_round(d, [68, 970, 340, 1040], "#f1f5f9", None, 16)
    d.text((98, 992), "定位：已获取", fill="#334155", font=f)
    draw_round(d, [370, 970, 712, 1040], "#f1f5f9", None, 16)
    d.text((400, 992), "照片：可选", fill="#334155", font=f)
    draw_round(d, [68, 1110, 712, 1190], "#1d4ed8", None, 20)
    d.text((296, 1133), "提交打卡", fill="white", font=f_b)
    d.text((48, 1270), "最近记录：06-22 查寝 · 教师待审核 · 定位正常", fill="#334155", font=f)
    out = ASSET_DIR / "prototype_student.png"
    img.save(out)
    images["student"] = out

    # Teacher/admin wide prototype
    def dashboard(kind: str, title: str, accent: str):
        im = Image.new("RGB", (1400, 850), "#f8fafc")
        dr = ImageDraw.Draw(im)
        dr.rectangle([0, 0, 1400, 88], fill=accent)
        dr.text((44, 28), title, fill="white", font=f_title)
        if kind == "teacher":
            draw_round(dr, [34, 120, 1366, 205], "#ffffff", "#e2e8f0")
            dr.text((60, 150), "筛选：软件2401  |  2026-06-22  |  查寝  |  待审核", fill="#334155", font=f)
            draw_round(dr, [34, 240, 850, 790], "#ffffff", "#e2e8f0")
            dr.text((70, 275), "待审核队列", fill="#0f172a", font=f_h)
            rows = [("李明 · 查寝打卡", "定位正常", "#f0fdf4"), ("周然 · 实习打卡", "定位异常", "#fff7ed"), ("陈晨 · 上课打卡", "定位正常", "#ffffff")]
            for i, row in enumerate(rows):
                y = 330 + i * 120
                draw_round(dr, [70, y, 810, y + 86], row[2], "#d1d5db")
                dr.text((100, y + 18), row[0], fill="#0f172a", font=f_b)
                dr.text((100, y + 52), "内容摘要：今日完成打卡并提交心得...", fill="#64748b", font=f)
                dr.text((660, y + 30), row[1], fill="#166534" if "正常" in row[1] else "#c2410c", font=f)
            draw_round(dr, [890, 240, 1366, 790], "#ffffff", "#e2e8f0")
            dr.text((930, 275), "打卡详情", fill="#0f172a", font=f_h)
            for idx, line in enumerate(["学生：李明 · 软件2401", "类型：查寝打卡", "时间：22:14", "定位：宿舍楼附近 · 正常"]):
                dr.text((930, 330 + idx * 42), line, fill="#475569", font=f)
            draw_round(dr, [930, 520, 1110, 585], "#16a34a", None)
            dr.text((990, 540), "通过", fill="white", font=f_b)
            draw_round(dr, [1140, 520, 1320, 585], "#fee2e2", "#fca5a5")
            dr.text((1200, 540), "驳回", fill="#b91c1c", font=f_b)
        else:
            labels = [("今日应打卡", "128"), ("已打卡", "109"), ("打卡率", "85%"), ("待审核/异常", "23")]
            for i, (a, b) in enumerate(labels):
                x = 34 + i * 340
                draw_round(dr, [x, 130, x + 300, 260], "#ffffff", "#e2e8f0")
                dr.text((x + 30, 160), a, fill="#64748b", font=f)
                dr.text((x + 30, 200), b, fill="#1d4ed8" if i == 2 else "#0f172a", font=f_title)
            draw_round(dr, [34, 310, 880, 760], "#ffffff", "#e2e8f0")
            dr.text((70, 350), "近 7 日打卡率趋势", fill="#0f172a", font=f_h)
            pts = [(90, 690), (210, 620), (330, 640), (450, 560), (570, 520), (700, 470), (830, 430)]
            dr.line(pts, fill="#1d4ed8", width=8)
            draw_round(dr, [930, 310, 1366, 760], "#ffffff", "#e2e8f0")
            dr.text((970, 350), "异常待处理", fill="#0f172a", font=f_h)
            for i, line in enumerate(["软件2401 · 定位异常 4 条", "未打卡学生 19 人", "查寝场景集中异常"]):
                y = 410 + i * 95
                draw_round(dr, [970, y, 1320, y + 65], "#fff7ed", "#fed7aa")
                dr.text((996, y + 20), line, fill="#334155", font=f)
        o = ASSET_DIR / f"prototype_{kind}.png"
        im.save(o)
        return o

    images["teacher"] = dashboard("teacher", "ZeKin 教师审核工作台", "#166534")
    images["admin"] = dashboard("admin", "ZeKin 管理统计概览", "#0f172a")
    return images


def set_cell_text(cell, text, bold=False, size=10.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(size)
    r.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))
    tc_w.set(qn("w:type"), "dxa")


def add_table(doc, caption, headers, rows, widths=None):
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
        if widths:
            set_cell_width(table.rows[0].cells[i], widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), size=10.5)
            if widths:
                set_cell_width(cells[i], widths[i])
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1, page_break=True):
    if level == 1:
        if page_break:
            doc.add_page_break()
        p = doc.add_paragraph(style="Heading 1")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        p = doc.add_paragraph(style="Heading 2")
    else:
        p = doc.add_paragraph(style="Heading 3")
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(style="Body Text")
    p.paragraph_format.first_line_indent = Pt(21)
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Body Text")
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Pt(21)
        p.paragraph_format.hanging_indent = Pt(10.5)
        p.add_run("• " + item)


def add_code(doc, code):
    for line in code.splitlines():
        p = doc.add_paragraph(style="Code")
        p.add_run(line if line else " ")


def add_field(paragraph, field):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def patch_page_number_type(section, fmt="decimal", start=None):
    sect_pr = section._sectPr
    pg = sect_pr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sect_pr.append(pg)
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))


def setup_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(2.4)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    body = styles["Body Text"]
    body.font.name = "宋体"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    body.font.size = Pt(12)
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(0)
    body.paragraph_format.first_line_indent = Pt(21)

    for style_name, size, align in [
        ("Heading 1", 16, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", 15, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 14, WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        st = styles[style_name]
        st.font.name = "黑体"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.font.size = Pt(size)
        st.font.bold = False
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.alignment = align
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(0)

    cap = styles["Caption"]
    cap.font.name = "宋体"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cap.font.size = Pt(10.5)
    cap.font.italic = False
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing = 1.5
    cap.paragraph_format.space_after = Pt(0)

    code = styles.add_style("Code", 1)
    code.font.name = "Times New Roman"
    code.font.size = Pt(12)
    code.paragraph_format.line_spacing = 1.5
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)

    return doc


def set_header_footer(section, show_footer=True, page_field="PAGE", literal_footer=None):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.text = ""
    r = header.add_run("实训-知勤打卡系统-需求分析")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(9)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = ""
    if not show_footer:
        return
    if literal_footer:
        rr = footer.add_run(literal_footer)
        rr.font.name = "宋体"
        rr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        rr.font.size = Pt(9)
        return
    for text in ["第 "]:
        rr = footer.add_run(text)
        rr.font.name = "宋体"
        rr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        rr.font.size = Pt(9)
    add_field(footer, page_field)
    rr = footer.add_run(" 页 共 ")
    rr.font.name = "宋体"; rr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体"); rr.font.size = Pt(9)
    add_field(footer, "NUMPAGES")
    rr = footer.add_run(" 页")
    rr.font.name = "宋体"; rr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体"); rr.font.size = Pt(9)


def add_picture(doc, path: Path, caption: str, width_cm=14.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    c = doc.add_paragraph(style="Caption")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.add_run(caption)


def build():
    uml_images = render_plantuml()
    proto_images = make_prototype_images()

    doc = setup_document()
    for sec in doc.sections:
        set_header_footer(sec, show_footer=False)

    # Cover
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("实训-知勤打卡系统\n需求分析报告")
    r.font.name = "黑体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(24); r.bold = True
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = sub.add_run("基于 UML 的 72 小时 MVP 需求分析与建模")
    rr.font.name = "宋体"; rr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    rr.font.size = Pt(14)
    meta = [
        "项目名称：知勤打卡系统（ZeKin）",
        "文档版本：V2.1",
        "项目周期：3 天 × 24 小时 = 72 小时极限编程",
        "技术约束：Python + FastAPI + Vue + MySQL",
        "团队成员：王韩韵、赵耀、胡钊炫、华心仪",
        "编制日期：2026 年 6 月 22 日",
    ]
    for _ in range(5):
        doc.add_paragraph()
    for m in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(m)
        run.font.name = "宋体"; run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)

    # Abstract
    doc.add_section(WD_SECTION.NEW_PAGE)
    patch_page_number_type(doc.sections[-1], "lowerRoman", 1)
    set_header_footer(doc.sections[-1], show_footer=True, literal_footer="第 i 页 共 47 页")
    h = doc.add_paragraph(style="Heading 1"); h.alignment = WD_ALIGN_PARAGRAPH.CENTER; h.add_run("摘  要")
    add_body(doc, "本文档面向“实训-知勤打卡系统”72 小时 MVP 开发，采用 UML 需求分析方法，将高校学生查寝、上课、实习三类打卡场景抽象为角色、用例、业务流程、系统协作和分析类模型。文档以业务语言说明学校、教师、学生在日常管理中的痛点，以 UML 图形化方式明确系统边界、参与者、核心用例、活动流程、顺序协作、时间窗口和领域对象，为后续前后端分离开发提供统一规格。")
    add_body(doc, "72 小时 MVP 的目标不是交付完整平台，而是完成“学生打卡、教师审核、统计概览、数据追溯”的最小业务闭环。AI 审核、人脸识别、补签、推送、积分等能力作为扩展点预留，不进入极限开发的强制范围。")
    add_body(doc, "关键词：知勤打卡；UML；用例规约；活动图；顺序图；时序图；分析类图；MVP")

    # TOC page
    doc.add_section(WD_SECTION.NEW_PAGE)
    set_header_footer(doc.sections[-1], show_footer=False)
    h = doc.add_paragraph(style="Heading 1"); h.alignment = WD_ALIGN_PARAGRAPH.CENTER; h.add_run("目  录")
    toc = [
        ("1 引言", "1"),
        ("2 项目背景与业务目标", "3"),
        ("3 72 小时 MVP 范围定义", "5"),
        ("4 参与者与用例模型", "8"),
        ("5 核心用例规约", "11"),
        ("6 业务流程活动建模", "20"),
        ("7 系统交互顺序建模", "25"),
        ("8 打卡业务时序建模", "29"),
        ("9 分析类模型", "31"),
        ("10 原型设计", "35"),
        ("11 团队分工与极限开发计划", "39"),
        ("12 验收标准与风险控制", "42"),
        ("附录 PlantUML 源码", "45"),
    ]
    for name, page in toc:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(22)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.first_line_indent = None
        run = p.add_run(f"{name}{'.' * max(4, 48 - len(name))}{page}")
        run.font.name = "宋体"; run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)

    # Main body
    doc.add_section(WD_SECTION.NEW_PAGE)
    patch_page_number_type(doc.sections[-1], "decimal", 1)
    set_header_footer(doc.sections[-1], show_footer=True)

    add_heading(doc, "1 引言", 1, page_break=False)
    add_body(doc, "知勤打卡系统的需求分析必须先回答一个业务问题：学校为什么需要一个数字化打卡系统，而不是继续使用人工点名、纸质签到或微信群接龙。高校学生管理的核心矛盾在于管理对象分散、管理时间固定、管理结果需要追溯，而传统方式依赖人工经验，难以形成持续、准确、可分析的数据资产。")
    add_body(doc, "本文档以 UML 为建模语言，将业务需求转换为可沟通、可验证、可交付的模型资产。叙事部分使用业务语言说明业务目标和约束，图形部分使用 PlantUML 表达用例图、活动图、顺序图、时序图和分析类图，确保团队在 72 小时极限开发中围绕同一套规格并行工作。")

    add_heading(doc, "2 项目背景与业务目标", 1)
    add_heading(doc, "2.1 业务痛点", 2)
    add_table(doc, "表 2.1 高校学生日常管理痛点", ["场景", "现状", "业务影响"], [
        ["查寝", "辅导员夜间逐间宿舍点名，记录依赖人工", "耗时长，信息滞后，异常学生难以及时定位"],
        ["上课", "课前纸质签到或口头点名", "占用教学时间，代签风险高，考勤数据难以统计"],
        ["实习", "学生分散在不同城市与企业", "学校难以持续掌握学生安全状态和实习动态"],
    ], [2.5, 5.2, 6.3])
    add_heading(doc, "2.2 产品定位", 2)
    add_body(doc, "知勤打卡系统不是单纯的签到工具，而是面向高校学生日常管理的数字化闭环系统。签到工具只回答“到没到”，而知勤打卡系统需要进一步回答“是否在规定场景完成任务、内容是否可审核、异常是否可追踪、数据是否可统计”。")
    add_heading(doc, "2.3 业务目标", 2)
    add_bullets(doc, [
        "让学生在移动端快速完成查寝、上课、实习三类打卡。",
        "让教师在一个工作台内查看所管班级的打卡状态并完成审核。",
        "让管理员或教师通过统计概览掌握今日完成率、待审核数和异常数。",
        "为后续 AI 审核、人脸识别、补签审批、通知提醒等能力预留数据结构和流程扩展点。",
    ])

    add_heading(doc, "3 72 小时 MVP 范围定义", 1)
    add_heading(doc, "3.1 极限开发约束", 2)
    add_body(doc, "本项目开发周期为 3 天，即 72 小时极限编程。该时间条件决定了系统必须坚持最小闭环，而不是追求功能大全。进入 MVP 的功能必须直接服务于演示和验收闭环；任何外部依赖强、联调风险高、非核心路径的能力均应降级为扩展点。")
    add_table(doc, "表 3.1 MVP 功能范围", ["层级", "功能", "处理策略"], [
        ["Must", "登录与角色识别", "必须实现，用于区分学生、教师、管理员"],
        ["Must", "三类打卡提交", "必须实现，覆盖查寝、上课、实习核心场景"],
        ["Must", "教师审核", "必须实现，形成管理闭环"],
        ["Must", "统计概览", "必须实现，支撑验收演示"],
        ["Should", "定位记录", "记录经纬度和异常状态，不强制范围校验"],
        ["Could", "AI 审核、人脸识别、补签、推送", "预留字段或接口，不进入 72 小时强制范围"],
    ], [2.0, 5.0, 7.0])
    add_heading(doc, "3.2 MVP 成功定义", 2)
    add_body(doc, "72 小时后，系统应能够完成以下演示：学生登录并提交查寝打卡，系统阻止重复打卡；教师登录并筛选今日待审核记录，对学生打卡进行通过或驳回；学生刷新历史记录后看到教师审核结果；统计页面展示今日应打卡、已打卡、打卡率、待审核和异常数量。")

    add_heading(doc, "4 参与者与用例模型", 1)
    add_heading(doc, "4.1 参与者定义", 2)
    add_table(doc, "表 4.1 参与者说明", ["参与者", "业务目标", "权限边界"], [
        ["学生", "完成规定场景下的打卡，查看审核结果", "只能创建和查看本人打卡数据"],
        ["辅导员/教师", "管理所管班级打卡，处理待审核和异常", "只能查看和审核所管班级数据"],
        ["管理员", "查看全局统计，维护基础数据", "具备全局查看和配置权限"],
        ["定位服务", "提供打卡经纬度", "仅作为外部辅助服务"],
    ], [3, 5.6, 5.6])
    if "00_system_usecase" in uml_images:
        add_picture(doc, uml_images["00_system_usecase"], "图 4.1 知勤打卡系统用例图", 14.2)
    add_body(doc, "用例模型显示系统边界内的核心能力包括用户登录、提交打卡、查看今日状态与历史、教师审核打卡和查看统计概览。其中提交打卡包含采集定位和校验重复打卡，教师审核包含记录审核结果，统计概览可扩展进入异常审核处理。")

    add_heading(doc, "5 核心用例规约", 1)
    use_cases = [
        ("UC-01 用户登录", "用户以合法身份进入系统，并获得与角色匹配的功能入口。", ["用户打开登录页", "输入账号和密码", "系统校验账号和密码", "系统识别角色", "系统跳转对应工作台"], ["账号不存在或密码错误时统一提示账号或密码错误", "账号禁用时提示联系管理员"]),
        ("UC-02 提交打卡", "学生在查寝、上课、实习场景下提交可追溯的打卡记录。", ["学生进入今日打卡页", "选择打卡类型", "填写内容", "系统获取定位", "学生提交", "系统校验重复规则", "系统保存待审核记录"], ["定位失败时允许提交但标记异常", "重复打卡时拒绝提交"]),
        ("UC-03 查看今日状态与历史", "学生确认今日任务完成情况，并追踪教师审核结果。", ["学生进入首页", "系统查询三类状态", "系统展示状态与最近记录", "学生查看审核结果和评语"], ["无记录时显示待打卡", "记录为空时显示空状态"]),
        ("UC-04 教师审核打卡", "教师快速处理班级打卡记录，形成管理闭环。", ["教师进入审核工作台", "选择班级日期类型状态", "系统校验班级权限", "教师打开详情", "教师通过或驳回", "系统保存审核记录"], ["无权限时拒绝访问", "驳回原因为空时提示补充"]),
        ("UC-05 查看统计概览", "管理人员用最短时间掌握今日打卡完成度和异常情况。", ["用户进入统计页", "系统识别角色范围", "系统聚合指标", "页面展示指标卡片", "用户进入审核列表处理异常"], ["无权限时拒绝访问", "暂无数据时展示空状态"]),
    ]
    for idx, (name, goal, main, ext) in enumerate(use_cases, 1):
        add_heading(doc, f"5.{idx} {name}", 2)
        add_table(doc, f"表 5.{idx} {name} 用例摘要", ["项目", "内容"], [
            ["业务目标", goal],
            ["优先级", "Must"],
            ["前置条件", "用户已具备合法账号；除登录用例外，用户已登录"],
            ["后置条件", "系统状态或业务记录被正确更新"],
        ], [3.0, 10.8])
        add_body(doc, "主成功场景：" + "；".join(main) + "。")
        add_body(doc, "扩展场景：" + "；".join(ext) + "。")

    add_heading(doc, "6 业务流程活动建模", 1)
    for i, (key, title_text) in enumerate([
        ("01_submit_checkin_activity", "提交打卡活动图"),
        ("02_teacher_review_activity", "教师审核活动图"),
        ("03_statistics_activity", "统计概览活动图"),
    ], 1):
        add_heading(doc, f"6.{i} {title_text}", 2)
        if key in uml_images:
            add_picture(doc, uml_images[key], f"图 6.{i} {title_text}", 14.2)
        add_body(doc, f"{title_text}用于说明该业务过程中的参与者动作、系统判断、异常分支和业务状态变化。活动图将开发人员容易误解的口头流程转化为可检查的分支结构，尤其适合 72 小时开发前的范围冻结。")

    add_heading(doc, "7 系统交互顺序建模", 1)
    for i, (key, title_text) in enumerate([
        ("04_submit_checkin_sequence", "提交打卡顺序图"),
        ("05_teacher_review_sequence", "教师审核顺序图"),
    ], 1):
        add_heading(doc, f"7.{i} {title_text}", 2)
        if key in uml_images:
            add_picture(doc, uml_images[key], f"图 7.{i} {title_text}", 14.2)
        add_body(doc, f"{title_text}展示页面边界对象、控制对象、服务对象、仓储对象和数据库之间的消息传递顺序。该模型可直接指导前后端接口、服务层和数据库访问层的实现分工。")

    add_heading(doc, "8 打卡业务时序建模", 1)
    if "06_checkin_timing" in uml_images:
        add_picture(doc, uml_images["06_checkin_timing"], "图 8.1 三类打卡业务时序图", 14.2)
    add_body(doc, "时序图用于表达查寝、上课、实习三类打卡在一天内的业务窗口。72 小时 MVP 不实现复杂排课与时间引擎，但需要在需求层面明确不同场景的开放时间、关闭时间和教师集中审核时间。")

    add_heading(doc, "9 分析类模型", 1)
    if "07_analysis_class_diagram" in uml_images:
        add_picture(doc, uml_images["07_analysis_class_diagram"], "图 9.1 知勤打卡系统分析类图", 14.0)
    add_body(doc, "分析类模型采用 BCE 方法划分边界类、控制类和实体类。边界类面向用户界面和外部服务，控制类承载用例流程和权限判断，实体类保存长期业务数据。该模型能够让团队明确前端、后端和数据库的协作边界。")
    add_table(doc, "表 9.1 分析类职责说明", ["类别", "代表类", "职责"], [
        ["Boundary", "StudentCheckinPage、TeacherReviewPage、AdminDashboardPage", "接收用户输入，展示业务状态，调用后端接口"],
        ["Control", "CheckinController、ReviewController、StatisticsController", "编排用例流程，校验业务规则，协调实体对象"],
        ["Entity", "User、Checkin、Review、ClassGroup、Course、LocationRule", "保存核心业务数据和对象关系"],
    ], [2.5, 5.5, 6.0])

    add_heading(doc, "10 原型设计", 1)
    add_heading(doc, "10.1 学生打卡页", 2)
    add_picture(doc, proto_images["student"], "图 10.1 学生打卡页原型样图", 8.2)
    add_body(doc, "学生端首屏应优先展示今日状态，帮助学生立即判断是否还有未完成任务。打卡表单保留类型、内容、定位、照片和提交按钮，避免在 72 小时 MVP 中引入复杂交互。")
    add_heading(doc, "10.2 教师审核页", 2)
    add_picture(doc, proto_images["teacher"], "图 10.2 教师审核页原型样图", 14.2)
    add_body(doc, "教师端的核心任务是处理待审核和异常记录，因此页面布局采用左侧队列、右侧详情的工作台结构，减少教师在列表和详情之间反复跳转的成本。")
    add_heading(doc, "10.3 管理统计页", 2)
    add_picture(doc, proto_images["admin"], "图 10.3 管理统计页原型样图", 14.2)
    add_body(doc, "管理统计页聚焦今日应打卡、已打卡、打卡率、待审核和异常数。趋势图在 72 小时 MVP 中可使用轻量图表或占位实现，优先保证统计口径准确。")

    add_heading(doc, "11 团队分工与极限开发计划", 1)
    add_table(doc, "表 11.1 72 小时团队分工", ["成员", "角色", "0-24 小时", "24-48 小时", "48-72 小时"], [
        ["王韩韵", "产品/QA", "冻结需求、用例、验收脚本", "跟进联调、记录缺陷", "组织四环境验收"],
        ["赵耀", "后端/数据库/部署", "FastAPI、MySQL、认证骨架", "打卡、审核、统计 API", "部署、修复、演示数据"],
        ["胡钊炫", "教师/管理端", "管理端骨架与登录", "审核列表、详情、统计页", "UI 修正、联调验收"],
        ["华心仪", "学生端", "学生端骨架与登录", "今日状态、打卡表单、历史", "移动端适配、联调验收"],
    ], [2.0, 2.3, 3.1, 3.4, 3.4])
    add_table(doc, "表 11.2 72 小时时间盒", ["时间", "目标", "交付物"], [
        ["H0-H6", "需求冻结", "用例、图、原型、验收脚本"],
        ["H6-H12", "项目骨架", "后端、前端、数据库最小启动"],
        ["H12-H24", "认证与基础数据", "登录、角色、演示数据"],
        ["H24-H36", "学生打卡闭环", "今日状态、提交打卡、历史"],
        ["H36-H48", "教师审核闭环", "审核列表、详情、通过/驳回"],
        ["H48-H60", "统计与联调", "指标概览、权限、异常处理"],
        ["H60-H72", "验收与修复", "四环境验收、演示脚本、缺陷修复"],
    ], [2.4, 4.2, 7.4])

    add_heading(doc, "12 验收标准与风险控制", 1)
    add_heading(doc, "12.1 四环境验收", 2)
    add_table(doc, "表 12.1 四环境验收标准", ["环境", "验收方式", "通过标准"], [
        ["本地开发环境", "成员本机启动前后端和数据库", "可完成登录、打卡、审核、统计"],
        ["服务器部署环境", "统一部署到服务器", "外部设备可访问演示地址"],
        ["用户验收环境", "使用演示账号走验收脚本", "业务方确认核心闭环可用"],
        ["生产环境测试", "使用生产配置冒烟测试", "不暴露调试信息，核心流程可完成"],
    ], [3.0, 5.2, 5.8])
    add_heading(doc, "12.2 风险控制", 2)
    add_bullets(doc, [
        "AI 审核接口不稳定：MVP 中降级为待审核状态，由教师审核兜底。",
        "定位不准确：MVP 只记录定位和异常标记，不强制地理围栏判断。",
        "前后端接口冲突：用用例规约和顺序图先冻结接口字段与状态枚举。",
        "72 小时联调时间不足：每 12 小时合并一次，任何超过 2 小时的问题必须降级或绕过。",
    ])

    add_heading(doc, "附  录", 1)
    add_heading(doc, "附录 1 PlantUML 源码", 2)
    for puml in sorted((ROOT / "diagrams").glob("*.puml")):
        add_heading(doc, f"附录 1.{len([p for p in doc.paragraphs if p.style.name == 'Heading 3']) + 1} {puml.name}", 3)
        add_code(doc, puml.read_text(encoding="utf-8"))
    add_heading(doc, "附录 2 参考资料", 2)
    refs = [
        "OMG UML 2.5.1：https://www.omg.org/spec/UML/2.5.1/About-UML",
        "OMG UML 资源页：https://www.omg.org/uml/",
        "PlantUML 官方文档：https://plantuml.com/",
        "PlantUML 顺序图语法：https://plantuml.com/sequence-diagram",
        "PlantUML 类图语法：https://plantuml.com/class-diagram",
        "Visual Paradigm 用例规约说明：https://www.visual-paradigm.com/guide/use-case/what-is-use-case-specification/",
    ]
    add_bullets(doc, refs)

    doc.core_properties.title = "实训-知勤打卡系统-需求分析"
    doc.core_properties.subject = "UML 需求分析与 72 小时 MVP 建模"
    doc.core_properties.author = "ZeKin 项目组"
    doc.save(OUT)
    print(OUT)
    print(f"rendered_uml={len(uml_images)}")


if __name__ == "__main__":
    build()
