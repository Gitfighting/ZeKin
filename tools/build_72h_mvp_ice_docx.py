from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "实训-知勤打卡系统-72小时MVP需求精简与里程碑报告.docx"


def add_field(paragraph, field):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, text, end])


def setup_page(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.4)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def set_header_footer(section, show_footer=True):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.text = ""
    r = header.add_run("实训-知勤打卡系统-72小时MVP")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(9)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = ""
    if not show_footer:
        return
    r = footer.add_run("第 ")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(9)
    add_field(footer, "PAGE")
    r = footer.add_run(" 页 共 ")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(9)
    add_field(footer, "NUMPAGES")
    r = footer.add_run(" 页")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(9)


def patch_page_start(section, start=1):
    sect_pr = section._sectPr
    pg = sect_pr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sect_pr.append(pg)
    pg.set(qn("w:start"), str(start))


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    body = styles["Body Text"]
    body.font.name = "宋体"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    body.font.size = Pt(12)
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.first_line_indent = Pt(21)
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(0)

    for style_name, size, align in [
        ("Heading 1", 16, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", 15, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 14, WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        st = styles[style_name]
        st.font.name = "黑体"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.font.size = Pt(size)
        st.font.bold = False
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.alignment = align
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(0)

    cap = styles["Caption"]
    cap.font.name = "宋体"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cap.font.size = Pt(10.5)
    cap.font.italic = False
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def h(doc, text, level=1, page_break=True):
    if level == 1:
        if page_break:
            doc.add_page_break()
        p = doc.add_paragraph(style="Heading 1")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        p = doc.add_paragraph(style="Heading 2")
    else:
        p = doc.add_paragraph(style="Heading 3")
    p.add_run(text)


def p(doc, text):
    para = doc.add_paragraph(style="Body Text")
    para.paragraph_format.first_line_indent = Pt(21)
    para.add_run(text)


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="Body Text")
        para.paragraph_format.first_line_indent = None
        para.paragraph_format.left_indent = Pt(21)
        para.paragraph_format.hanging_indent = Pt(10.5)
        para.add_run("• " + item)


def cell(cell, text, bold=False, size=10.5, align=None):
    cell.text = ""
    para = cell.paragraphs[0]
    if align:
        para.alignment = align
    run = para.add_run(str(text))
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade(cell_obj, fill):
    tc_pr = cell_obj._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def width(cell_obj, cm):
    tc_pr = cell_obj._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(cm).twips)))
    tc_w.set(qn("w:type"), "dxa")


def tbl(doc, caption, headers, rows, widths):
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for i, header in enumerate(headers):
        c = table.rows[0].cells[i]
        cell(c, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(c, "E8EEF5")
        width(c, widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell(cells[i], value)
            width(cells[i], widths[i])
    doc.add_paragraph()


def build():
    doc = Document()
    setup_page(doc.sections[0])
    setup_styles(doc)
    set_header_footer(doc.sections[0], show_footer=False)

    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("实训-知勤打卡系统\n72小时MVP需求精简与里程碑报告")
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(23)
    r.bold = True
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("ICE 优先级法 · 最简技术架构 · 三日交付计划")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(14)
    for _ in range(5):
        doc.add_paragraph()
    for line in [
        "项目名称：知勤打卡系统（AI 思政打卡 MVP）",
        "文档版本：V1.1",
        "项目周期：72 小时极限 MVP",
        "技术约束：Python + FastAPI + Vue + MySQL",
        "编制日期：2026 年 6 月 23 日",
    ]:
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = meta.add_run(line)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(10.5)

    doc.add_section(WD_SECTION.NEW_PAGE)
    setup_page(doc.sections[-1])
    set_header_footer(doc.sections[-1], show_footer=True)
    patch_page_start(doc.sections[-1], 1)

    h(doc, "1 一句话说明这个项目", 1, page_break=False)
    p(doc, "知勤打卡系统可以理解成“学校版打卡小程序”。学生每天在手机上完成查寝、上课或实习打卡；老师在后台看谁打了、谁没打、谁需要审核；管理员看整体数据。")
    p(doc, "72 小时 MVP 的目标不是做一个完美产品，而是先把最核心的业务流程跑通。就像先搭一座能走人的小桥，再慢慢加栏杆、灯光和装饰。")

    h(doc, "2 需求精简：用 ICE 优先级法决定先做什么", 1)
    h(doc, "2.1 功能剪减思路", 2)
    p(doc, "ICE 是一个很简单的打分方法。Impact 表示这个功能对用户有没有用，Confidence 表示 72 小时内做出来的把握有多大，Ease 表示实现是否容易。三个分数加起来越高，就越应该先进 MVP。")
    tbl(doc, "表 2.1 ICE 功能优先级评分", ["功能模块", "I 影响力", "C 可行性", "E 容易度", "总分", "MVP 决策"], [
        ["用户登录/注册", "9", "9", "9", "27", "必做（P0）"],
        ["学生提交打卡", "10", "8", "7", "25", "必做（P0）"],
        ["打卡历史查看", "7", "9", "8", "24", "必做（P0）"],
        ["教师审核面板", "9", "7", "6", "22", "简化做（P1）"],
        ["统计看板", "6", "7", "5", "18", "简化做（P1）"],
        ["AI 内容审核", "5", "4", "3", "12", "不做（P2）"],
        ["人脸识别", "3", "3", "2", "8", "不做（P2）"],
        ["语音打卡", "2", "2", "2", "6", "不做（P2）"],
        ["积分/排行榜", "4", "3", "3", "10", "不做（P2）"],
    ], [3.8, 1.8, 1.8, 1.8, 1.7, 3.3])

    h(doc, "2.2 MVP 剪减后的最终功能清单", 2)
    tbl(doc, "表 2.2 72 小时 MVP 最终功能清单", ["序号", "功能模块", "负责人", "功能说明"], [
        ["1", "用户登录/注册", "赵耀", "手机号注册 + JWT 认证，支持学生/教师/管理员三种角色"],
        ["2", "学生提交打卡", "王韩韵", "文字心得 + 照片上传 + GPS 定位，支持查寝/上课/实习三种类型"],
        ["3", "打卡历史查看", "胡钊炫", "学生查看自己的打卡记录、审核状态和筛选"],
        ["4", "教师审核面板", "华心仪", "查看班级学生打卡列表，通过/拒绝并留评语（简化版）"],
        ["5", "简易统计看板", "华心仪", "展示打卡率、未打卡人数、连续打卡天数"],
    ], [1.4, 3.2, 2.0, 8.6])

    h(doc, "2.3 为什么剪掉 AI 审核", 2)
    bullets(doc, [
        "对外部 API 依赖太重：72 小时内如果 API 出问题，整个打卡流程可能卡死。",
        "审核结果不确定：AI 可能误判，学生和老师都需要额外解释，我们没有时间调优。",
        "教师审核已足够：MVP 阶段只要教师手动审核能走通，就能证明核心流程成立。",
        "设计原则：MVP 是验证核心流程可行，不是一次性做完美产品。AI 审核放到 V1.1 版本。",
    ])

    h(doc, "3 技术架构设计：最简原则", 1)
    h(doc, "3.1 分层架构", 2)
    tbl(doc, "表 3.1 最简技术架构", ["层级", "负责", "技术选型"], [
        ["前端层", "用户看到的页面、按钮、列表", "Vue 3 + Element Plus + Vite"],
        ["后端 API 层", "接收前端请求，处理登录、打卡、审核等业务", "FastAPI + Python 3.11"],
        ["数据存储层", "保存用户、打卡记录、审核记录", "MySQL 8.0"],
        ["部署层", "把项目托管到服务器，让别人能访问", "Docker Compose 一键部署"],
    ], [2.8, 6.6, 5.0])
    p(doc, "高中生可以这样理解：Vue 是门面，FastAPI 是服务台，MySQL 是档案柜，Docker Compose 是一键启动按钮。每一层只做自己的事，系统就不容易乱。")

    h(doc, "3.2 数据库设计：仅 3 张表", 2)
    tbl(doc, "表 3.2 MVP 最小数据库表", ["表名", "存什么", "关键字段"], [
        ["users", "用户信息", "id, username, password_hash, real_name, role(student/teacher/admin), class_name, phone"],
        ["checkins", "打卡记录", "id, user_id, type(dorm/class/internship), content, photo_url, lat, lng, status(pending/approved/rejected), created_at"],
        ["reviews", "审核记录", "id, checkin_id, reviewer_id, action(approved/rejected), comment, created_at"],
    ], [2.4, 3.2, 8.8])
    p(doc, "原版需求里的 courses、allowed_locations、makeup_requests 暂时剪掉。原因很简单：课程、定位白名单、补签都很有用，但它们不会影响“学生打卡 → 教师审核 → 统计展示”这条主路。")

    h(doc, "4 72 小时里程碑设计", 1)
    p(doc, "团队 4 人，72 小时按 3 天推进。真实开发中大约三分之一时间会花在沟通、联调、修问题上，所以计划必须留出对齐时间。")
    h(doc, "4.1 Day 1：基础设施 + 用户登录", 2)
    tbl(doc, "表 4.1 Day 1 任务安排", ["时间段", "负责人", "任务内容", "交付物"], [
        ["09:00-10:30", "全体", "项目启动会：确认技术选型、环境搭建、任务分拆", "GitHub 仓库初始化、Docker Compose 配置"],
        ["10:30-12:00", "赵耀", "后端：搭建 FastAPI 脚手架，MySQL 连接池配置", "API 服务启动成功"],
        ["10:30-12:00", "王韩韵", "前端：Vue 3 + Vite 项目初始化，Element Plus 集成", "前端项目脚手架运行"],
        ["10:30-12:00", "胡钊炫 + 华心仪", "数据库设计确认，创建 users/checkins/reviews 表", "数据库表创建完成"],
        ["14:00-17:00", "赵耀", "开发用户注册/登录 API（JWT 生成与验证）", "登录接口通过 Postman 测试"],
        ["14:00-17:00", "王韩韵", "开发登录页面 + 注册页面（表单验证、JWT 存储）", "用户能正常注册并登录"],
        ["14:00-17:00", "胡钊炫 + 华心仪", "后端基础 API 封装（响应模型、错误处理中间件）", "统一 API 响应格式"],
        ["19:00-21:00", "全体", "进度对齐：沟通接口联调问题，确认明天任务", "第一天进度报告"],
    ], [2.4, 3.0, 6.2, 5.0])
    bullets(doc, ["用户能注册并登录系统", "JWT Token 正常生成和验证", "MySQL 三张表创建完成", "前后端项目脚手架可运行"])

    h(doc, "4.2 Day 2：核心功能开发", 2)
    tbl(doc, "表 4.2 Day 2 任务安排", ["时间段", "负责人", "任务内容", "交付物"], [
        ["09:00-12:00", "王韩韵", "学生端：打卡页面开发（类型切换、文字输入、照片上传、GPS 获取）", "学生能提交打卡"],
        ["09:00-12:00", "赵耀", "后端：打卡提交 API（接收内容+照片+定位）", "打卡 API 可正常接收数据"],
        ["09:00-12:00", "胡钊炫", "后端：打卡历史查询 API（按用户/日期筛选）", "打卡列表接口完成"],
        ["09:00-12:00", "华心仪", "前端：教师审核列表页面（数据展示、操作按钮）", "审核列表页面架构"],
        ["14:00-17:00", "赵耀 + 胡钊炫", "后端：教师审核 API（批量查询班级打卡 + 通过/拒绝）", "审核接口完成"],
        ["14:00-17:00", "王韩韵 + 华心仪", "前端：学生打卡历史页面 + 教师审核页面联调", "前端页面联调通过"],
        ["19:00-21:00", "全体", "联调测试：模拟学生打卡 → 教师审核的完整流程", "核心业务流程走通"],
    ], [2.4, 3.0, 6.2, 5.0])
    bullets(doc, ["学生能提交打卡（含照片和定位）", "学生能查看自己的打卡历史", "教师能在审核面板查看班级打卡并操作", "前端与后端 API 联调通过"])

    h(doc, "4.3 Day 3：统计 + 优化 + 部署", 2)
    tbl(doc, "表 4.3 Day 3 任务安排", ["时间段", "负责人", "任务内容", "交付物"], [
        ["09:00-12:00", "华心仪", "前端：统计看板页面（打卡率、未打卡人数、连续打卡天数图表）", "统计看板可见"],
        ["09:00-12:00", "赵耀", "后端：统计 API 开发（打卡率聚合查询、排序）", "统计接口完成"],
        ["09:00-12:00", "王韩韵", "前端：页面 UI 美化、响应式适配、空状态提示", "页面视觉完善"],
        ["09:00-12:00", "胡钊炫", "后端：全局异常处理、请求日志、参数校验加固", "系统稳定性提升"],
        ["14:00-17:00", "全体", "Docker Compose 一键部署，整合前端静态资源 + 后端 API", "项目部署到服务器"],
        ["17:00-19:00", "全体", "上线前测试：完整流程测试、跨浏览器兼容性检查", "功能测试通过"],
        ["19:00-21:00", "全体", "交付准备：编写 README、录屏演示视频、部署文档", "MVP 交付包"],
    ], [2.4, 3.0, 6.2, 5.0])
    bullets(doc, ["统计看板正常展示数据", "Docker Compose 一键启动整个项目", "72 小时限时内系统稳定运行", "有 README 和部署文档"])

    h(doc, "5 MVP API 接口清单：仅 8 个接口", 1)
    tbl(doc, "表 5.1 MVP API 接口清单", ["方法", "路径", "功能说明", "权限", "负责人"], [
        ["POST", "/api/auth/register", "用户注册", "公开", "赵耀"],
        ["POST", "/api/auth/login", "用户登录，返回 JWT", "公开", "赵耀"],
        ["GET", "/api/auth/me", "获取当前登录用户信息", "登录用户", "赵耀"],
        ["POST", "/api/checkins", "提交打卡（内容+照片+定位）", "学生", "王韩韵"],
        ["GET", "/api/checkins/my", "获取我的打卡列表", "学生", "胡钊炫"],
        ["GET", "/api/teacher/checkins", "获取班级打卡列表", "教师", "华心仪"],
        ["POST", "/api/teacher/reviews", "提交审核结果（通过/拒绝）", "教师", "华心仪"],
        ["GET", "/api/stats/overview", "获取统计概览（打卡率等）", "教师/管理员", "赵耀"],
    ], [1.8, 4.6, 4.7, 2.2, 2.2])

    h(doc, "6 团队任务分配总览", 1)
    tbl(doc, "表 6.1 团队三日任务总览", ["团队成员", "Day 1 任务", "Day 2 任务", "Day 3 任务"], [
        ["王韩韵", "前端脚手架、登录/注册页面", "学生打卡页面、历史页面", "UI 美化、响应式适配"],
        ["赵耀", "后端脚手架、注册/登录 API", "打卡提交 API、审核 API", "统计 API、Docker 部署"],
        ["胡钊炫", "数据库表设计、基础 API 封装", "打卡历史查询 API、联调", "全局异常处理、参数校验"],
        ["华心仪", "数据库创建、响应模型封装", "教师审核页面、联调", "统计看板页面、图表"],
    ], [2.4, 4.0, 4.0, 4.0])
    h(doc, "6.1 协作规范", 2)
    bullets(doc, [
        "每天 09:00 和 21:00 各开一次 15 分钟进度对齐会。",
        "代码提交前必须自测，后端接口优先用 Postman 测试。",
        "每个功能开一个 feature 分支，完成后合并到 dev 分支。",
        "关键接口变更必须在群里通知所有人。",
        "遇到阻塞超过 1 小时就沟通，不要一个人硬扛。",
    ])

    h(doc, "7 风险评估与应对策略", 1)
    tbl(doc, "表 7.1 风险与应对", ["风险事项", "可能影响", "应对措施"], [
        ["前后端接口数据格式不一致", "联调失败，前端无法显示后端数据", "提前定义 API 响应规范，用 Postman 先测试后端"],
        ["代码冲突", "合并代码时出现冲突，修复浪费时间", "每人负责不同模块，尽量避免修改同一个文件"],
        ["环境不一致问题", "本地能跑，服务器跑不起来", "使用 Docker Compose 统一环境，不在本地直接堆依赖"],
        ["进度延迟", "某个功能花的时间超预期", "准备降级方案：统计看板可以用简单列表代替图表"],
        ["数据库性能", "查询慢影响用户体验", "为常用查询字段建索引，分页查询不过度加载"],
    ], [4.0, 5.0, 5.0])

    h(doc, "8 MVP 非功能需求：简化版", 1)
    tbl(doc, "表 8.1 非功能需求", ["类别", "MVP 要求", "具体指标"], [
        ["性能", "API 响应速度", "API 响应 < 1s（MVP 放宽要求）"],
        ["安全", "密码存储 + 身份认证", "密码 bcrypt 哈希 + JWT Token，禁用明文"],
        ["兼容性", "主流浏览器", "Chrome / Edge / Firefox 最新版（MVP 不测试 IE）"],
        ["可用性", "基本可用", "Docker 自动重启，MVP 不要求 99.9% 在线率"],
    ], [3.0, 4.5, 6.5])

    h(doc, "9 MVP 之后的 V1.1 路线图", 1)
    tbl(doc, "表 9.1 V1.1 候选功能", ["功能", "价值", "预计工作量"], [
        ["AI 内容审核", "自动判断打卡内容是否合规，减轻教师审核量", "2-3 天"],
        ["定位验证（GPS 固定范围）", "防止学生远程代打卡，提升可信度", "1-2 天"],
        ["补签申请流程", "学生忘记打卡时有正式补签通道", "1-2 天"],
        ["数据导出（Excel）", "教师月底上报需要导出考勤数据", "0.5-1 天"],
        ["打卡提醒通知", "减少忘记打卡的情况", "1 天"],
        ["人脸识别签到", "防止代签，最高安全级别", "3-5 天"],
    ], [4.0, 7.0, 3.0])

    h(doc, "10 参考资料", 1)
    bullets(doc, [
        "FastAPI 官方文档 Bigger Applications：APIRouter 用于组织大型应用的多文件路由。https://fastapi.tiangolo.com/tutorial/bigger-applications/",
        "Vue 官方文档 Composables：composable 用来封装和复用有状态逻辑。https://vuejs.org/guide/reusability/composables",
        "MySQL 8.4 官方文档 Foreign Key Constraints：InnoDB 支持外键约束并用于维护引用完整性。https://dev.mysql.com/doc/refman/8.4/en/create-table-foreign-keys.html",
        "GitHub Docs Protected Branches：分支保护可限制删除、强推，并要求状态检查或 Pull Request。https://docs.github.com/repositories/configuring-branches-and-merges-in-your-repository/managing-protected-branches/about-protected-branches",
        "GitHub Docs Pull Request Reviews：PR Review 可以帮助团队在合并前评论、建议和审批代码。https://docs.github.com/pull-requests/collaborating-with-pull-requests/reviewing-changes-in-pull-requests/about-pull-request-reviews",
    ])

    doc.core_properties.title = "实训-知勤打卡系统-72小时MVP需求精简与里程碑报告"
    doc.core_properties.author = "ZeKin 项目组"
    doc.core_properties.subject = "ICE 优先级法、72小时MVP、里程碑设计"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
